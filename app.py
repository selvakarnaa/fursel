"""
app.py — SAESL Subtask Lookup & Solumina Paste Tool

Images appear at their EXACT position in the subtask content flow
(same order as in the PDF). Paste to Solumina sends rich HTML that
includes both text and inline images.

Word mode: builds a .docx with all content + figures embedded,
           then opens it in WordPad automatically.
"""

import argparse
import base64
import mimetypes
import os
import re
import subprocess
import tempfile
from pathlib import Path

import pyperclip
from flask import Flask, render_template_string, request, jsonify, send_file, abort

from database import Database
from formatter import format_for_clipboard

# ─────────────────────────────────────────────
# OUTPUT MODE — change this single line to switch behaviour
# ─────────────────────────────────────────────
# Options:
#   "solumina"  — rich HTML to Windows clipboard via CF_HTML, auto Ctrl+V into Solumina
#   "word"      — builds .docx with embedded images and opens in WordPad
#   "notepad"   — plain text copied; paste with Ctrl+V into Notepad (or any plain-text app)
OUTPUT_MODE = "word"   # ← change only this line

app  = Flask(__name__)
DB_PATH = "output/saesl.db"

# ─────────────────────────────────────────────
# Rich HTML builder (used for both UI and paste)
# ─────────────────────────────────────────────

def build_rich_html(data: dict, base_url: str = "") -> str:
    """
    Build a self-contained HTML block for the subtask with images
    embedded at their exact content-flow position.
    When base_url is empty, images are base64-embedded (for clipboard paste).
    When base_url is set, images use <img src="…/api/image?path=…"> (for UI).
    """
    lines = []

    def img_src(path: str) -> str:
        if base_url:
            return f"{base_url}/api/image?path={path}"
        try:
            mime, _ = mimetypes.guess_type(path)
            data_b64 = base64.b64encode(Path(path).read_bytes()).decode()
            return f"data:{mime or 'image/png'};base64,{data_b64}"
        except Exception:
            return ""

    for ci in data.get("content_items", []):
        t = ci.get("type", "")

        if t == "image":
            src = img_src(ci["path"])
            if src:
                raw_label = ci.get("label", "")
                disp_label = re.sub(r"__p\d+__img\d+\.[^.]+$", "", raw_label).replace("_", "-") or raw_label
                lines.append(
                    f'<div style="margin:12px 0;text-align:center">'
                    f'<img src="{src}" alt="{disp_label}" '
                    f'style="max-width:100%;border:1px solid #dde3ef;border-radius:6px;padding:4px">'
                    f'</div>'
                )

        elif t == "callout":
            text = ci.get("text", "")
            low  = text.lower()
            if low.startswith("caution"):
                style = "border-left:4px solid #e8a020;background:#fef9ec;color:#7d5a00"
                icon  = "⚠️"
            elif low.startswith("warning") or low.startswith("notice"):
                style = "border-left:4px solid #c0392b;background:#fdecea;color:#7b1a14"
                icon  = "🚨"
            else:
                style = "border-left:4px solid #3498db;background:#eaf4fd;color:#1a5276"
                icon  = "ℹ️"
            lines.append(
                f'<div style="{style};padding:10px 14px;margin:8px 0;border-radius:4px;'
                f'font-size:13px;line-height:1.6">{icon} {_esc(text)}</div>'
            )

        elif t in ("step", "text", "row"):
            text = ci.get("text", "")
            clean = re.sub(r"^\s*(?:[A-Z]\.|\(\d+\)|\([a-z]\)|\d{1,2}\.)\s+", "", text)
            lines.append(
                f'<p style="font-family:Courier New,monospace;font-size:13px;'
                f'color:#2c3e50;margin:6px 0;line-height:1.65">{_esc(clean)}</p>'
            )

    return "\n".join(lines)


def _esc(s: str) -> str:
    return (s.replace("&", "&amp;").replace("<", "&lt;")
             .replace(">", "&gt;").replace('"', "&quot;"))


# ─────────────────────────────────────────────
# DOCX builder — for Word/WordPad mode
# ─────────────────────────────────────────────

def build_docx(data: dict) -> bytes:
    """
    Build a .docx document from subtask data with:
    - Header block (subtask ID, task, revision date)
    - Content flow: callouts, steps/text, and INLINE images at exact positions
    - Figure references section (figure images embedded directly)
    - Data cards section
    NO cross-references section.
    Returns raw .docx bytes.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Helper: add heading ──────────────────────────────────────────
    def add_heading(text, level=1, color=None):
        p = doc.add_heading(text, level=level)
        if color:
            for run in p.runs:
                run.font.color.rgb = RGBColor(*color)
        return p

    # ── Helper: add mono paragraph ────────────────────────────────────
    def add_mono(text, bold=False, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        p.paragraph_format.space_after = Pt(2)
        return p

    # ── Helper: add callout box (indented paragraph with label) ───────
    def add_callout(text):
        low = text.lower()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        if low.startswith("caution"):
            label_color = RGBColor(0x7d, 0x5a, 0x00)
            icon = "⚠ "
        elif low.startswith("warning") or low.startswith("notice"):
            label_color = RGBColor(0x7b, 0x1a, 0x14)
            icon = "🚨 "
        else:
            label_color = RGBColor(0x1a, 0x52, 0x76)
            icon = "ℹ "
        run = p.add_run(icon + text)
        run.font.name = "Segoe UI"
        run.font.size = Pt(10)
        run.font.color.rgb = label_color
        return p

    # ── Helper: embed image from path ─────────────────────────────────
    def add_image_from_path(path_str, label=""):
        p_path = Path(path_str)
        if not p_path.exists():
            return
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(p_path), width=Inches(5.5))
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            if label:
                disp = re.sub(r"__p\d+__img\d+\.[^.]+$", "", label).replace("_", "-") or label
                cap = doc.add_paragraph(disp)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap_run = cap.runs[0]
                cap_run.font.size = Pt(8)
                cap_run.font.color.rgb = RGBColor(0x6b, 0x7a, 0x99)
                cap_run.italic = True
                cap.paragraph_format.space_after = Pt(6)
        except Exception:
            pass

    # ── REFERENCE HEADER ─────────────────────────────────────────────
    title_p = doc.add_heading("", level=1)
    title_p.clear()
    run = title_p.add_run(data.get("subtask_id", ""))
    run.font.name  = "Courier New"
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x0f, 0x2a, 0x4a)

    if data.get("title"):
        sub_p = doc.add_paragraph(data["title"])
        sub_p.runs[0].italic = True
        sub_p.runs[0].font.color.rgb = RGBColor(0x6b, 0x7a, 0x99)
        sub_p.paragraph_format.space_after = Pt(2)

    meta_lines = []
    if data.get("task_id"):
        meta_lines.append(f"TASK          : {data['task_id']}")
    if data.get("revision_date"):
        meta_lines.append(f"REVISION DATE : {data['revision_date']}")
    for ml in meta_lines:
        add_mono(ml)

    doc.add_paragraph()  # spacer

    # ── CONTENT FLOW (steps, callouts, inline images) ─────────────────
    step_num = 0
    in_proc  = False

    for ci in data.get("content_items", []):
        t = ci.get("type", "")

        if t == "image":
            add_image_from_path(ci.get("path", ""), ci.get("label", ""))

        elif t == "callout":
            add_callout(ci.get("text", ""))

        elif t in ("step", "text", "row"):
            if not in_proc:
                h = doc.add_heading("PROCEDURE", level=2)
                h.runs[0].font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)
                in_proc = True
            step_num += 1
            raw = ci.get("text", "")
            clean = re.sub(r"^\s*(?:[A-Z]\.|\(\d+\)|\([a-z]\)|\d{1,2}\.)\s+", "", raw)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            num_run = p.add_run(f"{step_num:02d}.  ")
            num_run.font.name  = "Courier New"
            num_run.font.size  = Pt(10)
            num_run.font.bold  = True
            num_run.font.color.rgb = RGBColor(0x29, 0x52, 0xa3)
            txt_run = p.add_run(clean)
            txt_run.font.name = "Courier New"
            txt_run.font.size = Pt(10)

    # ── FIGURE REFERENCES with embedded images ────────────────────────
    fig_refs = data.get("figure_refs", [])
    if fig_refs:
        doc.add_paragraph()
        h = doc.add_heading("FIGURE REFERENCES", level=2)
        h.runs[0].font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)

        # Open DB to resolve figure images
        try:
            db = Database(DB_PATH)
            for fig_id in fig_refs:
                fig = db.get_figure_image(fig_id)
                cap_text = f"Fig {fig_id}"
                if fig:
                    if fig.get("caption"):
                        cap_text = fig["caption"]
                    img_path = fig.get("image_path", "")
                    if img_path and Path(img_path).exists():
                        lbl_p = doc.add_paragraph()
                        lbl_run = lbl_p.add_run(f"Figure {fig_id}")
                        lbl_run.font.bold = True
                        lbl_run.font.size = Pt(10)
                        lbl_run.font.color.rgb = RGBColor(0x0f, 0x2a, 0x4a)
                        add_image_from_path(img_path, cap_text)
                        continue
                # No image found — just text reference
                p = doc.add_paragraph()
                run = p.add_run(f"  Refer to Fig {fig_id}")
                run.font.name = "Courier New"
                run.font.size = Pt(10)
            db.close()
        except Exception:
            for fig_id in fig_refs:
                p = doc.add_paragraph()
                run = p.add_run(f"  Refer to Fig {fig_id}")
                run.font.name = "Courier New"
                run.font.size = Pt(10)

    # ── DATA CARDS ────────────────────────────────────────────────────
    dcs = data.get("data_cards", [])
    if dcs:
        doc.add_paragraph()
        h = doc.add_heading("ADDITIONAL DOCUMENTS", level=2)
        h.runs[0].font.color.rgb = RGBColor(0x1a, 0x3a, 0x6b)
        for dc in dcs:
            p = doc.add_paragraph()
            run = p.add_run(f"DATA CARD : {dc}")
            run.font.name = "Courier New"
            run.font.size = Pt(10)

    # ── Serialize ─────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────
# HTML UI Template
# ─────────────────────────────────────────────

HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>SAESL Subtask Lookup</title>
<style>
:root {
  --navy:#0f2a4a; --navy2:#1a3a6b; --navy3:#2952a3;
  --accent:#e8a020; --green:#1e7e34; --red:#c0392b;
  --bg:#f0f3f8; --card:#ffffff; --border:#dde3ef;
  --text:#1c2535; --muted:#6b7a99;
  --mono:'Courier New','Lucida Console',monospace;
}
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Segoe UI',Arial,sans-serif;background:var(--bg);color:var(--text);min-height:100vh}

header{
  background:linear-gradient(135deg,var(--navy) 0%,var(--navy2) 100%);
  color:white;padding:0 36px;height:64px;display:flex;align-items:center;gap:20px;
  box-shadow:0 2px 12px rgba(0,0,0,.25);position:sticky;top:0;z-index:100
}
.hlogo{width:36px;height:36px;background:var(--accent);border-radius:8px;
  display:flex;align-items:center;justify-content:center;
  font-size:18px;font-weight:800;color:var(--navy);flex-shrink:0}
.htitle{font-size:17px;font-weight:700;letter-spacing:.3px}
.hsub{font-size:12px;opacity:.65;margin-top:2px}
.hstats{margin-left:auto;display:flex;gap:20px}
.hstat .val{font-size:18px;font-weight:700;color:var(--accent)}
.hstat .lbl{font-size:10px;opacity:.6;text-transform:uppercase;letter-spacing:.6px}

.wrap{max-width:860px;margin:0 auto;padding:32px 24px}

.search-card{background:var(--card);border-radius:14px;padding:28px 32px;
  box-shadow:0 2px 16px rgba(15,42,74,.08);margin-bottom:28px;border:1px solid var(--border)}
.slabel{font-size:11px;font-weight:700;color:var(--muted);text-transform:uppercase;
  letter-spacing:1px;margin-bottom:12px;display:block}
.srow{display:flex;gap:12px;align-items:center}
.srow input{
  flex:1;padding:13px 18px;border:2px solid var(--border);border-radius:10px;
  font-size:15px;font-family:var(--mono);letter-spacing:.5px;color:var(--text);
  transition:border-color .18s,box-shadow .18s;background:#f8faff}
.srow input:focus{outline:none;border-color:var(--navy2);background:#fff;
  box-shadow:0 0 0 4px rgba(26,58,107,.08)}
.srow input::placeholder{color:#b0bad0}
.btn{padding:13px 24px;border:none;border-radius:10px;font-size:14px;font-weight:600;
  cursor:pointer;display:flex;align-items:center;gap:7px;transition:all .15s;white-space:nowrap}
.btn-primary{background:var(--navy2);color:white}
.btn-primary:hover{background:var(--navy);transform:translateY(-1px);box-shadow:0 4px 12px rgba(26,58,107,.3)}
.btn-success{background:var(--green);color:white}
.btn-success:hover{background:#166128;transform:translateY(-1px)}
.btn-success:disabled{background:#b0bac8;cursor:not-allowed;transform:none;box-shadow:none}
.btn-ghost{background:transparent;border:1.5px solid var(--border);color:var(--muted);
  padding:10px 16px;font-size:13px}
.btn-ghost:hover{border-color:var(--navy2);color:var(--navy2);background:#f0f4ff}
.status{padding:10px 16px;border-radius:8px;font-size:13px;margin-top:14px;
  display:none;align-items:center;gap:8px}
.status.error{background:#fdecea;color:var(--red);border:1px solid #fac8c3;display:flex}
.status.success{background:#eafaf1;color:var(--green);border:1px solid #b7e8c5;display:flex}
.status.info{background:#e8f0fd;color:var(--navy2);border:1px solid #c5d5f5;display:flex}

.empty-state{text-align:center;padding:80px 24px;color:var(--muted)}
.empty-icon{font-size:52px;margin-bottom:16px;opacity:.5}
.loader{display:flex;flex-direction:column;align-items:center;justify-content:center;
  padding:60px;gap:16px;color:var(--muted)}
.spinner{width:36px;height:36px;border:3px solid var(--border);
  border-top-color:var(--navy2);border-radius:50%;animation:spin .7s linear infinite}
@keyframes spin{to{transform:rotate(360deg)}}
@keyframes fadeIn{from{opacity:0;transform:translateY(8px)}to{opacity:1;transform:none}}

/* Result card */
.result-card{background:var(--card);border-radius:14px;
  box-shadow:0 2px 16px rgba(15,42,74,.08);border:1px solid var(--border);
  overflow:hidden;animation:fadeIn .25s ease}
.rh{background:linear-gradient(135deg,var(--navy) 0%,var(--navy2) 100%);
  padding:20px 28px;display:flex;align-items:flex-start;justify-content:space-between;gap:16px}
.rh-id{font-size:20px;font-weight:700;color:white;font-family:var(--mono);letter-spacing:.5px}
.rh-title{font-size:14px;color:rgba(255,255,255,.75);margin-top:4px;font-style:italic}
.rh-actions{display:flex;gap:10px;flex-shrink:0;align-items:center}

.stat-strip{display:flex;gap:0;border-bottom:1px solid var(--border);
  background:#f8faff;overflow-x:auto}
.stat-chip{padding:10px 20px;font-size:12px;color:var(--muted);
  border-right:1px solid var(--border);display:flex;align-items:center;gap:6px;white-space:nowrap}
.stat-chip strong{color:var(--text);font-weight:700}

/* Content flow — single column, images inline */
.content-flow{padding:28px 32px}

.section-head{font-size:11px;font-weight:700;text-transform:uppercase;
  letter-spacing:1px;color:var(--muted);margin:24px 0 10px;
  display:flex;align-items:center;gap:8px}
.section-head:first-child{margin-top:0}
.section-head::after{content:'';flex:1;height:1px;background:var(--border)}

.callout-box{border-left:4px solid;border-radius:4px;padding:10px 14px;
  margin-bottom:8px;font-size:13px;line-height:1.6}
.cb-note{border-color:#3498db;background:#eaf4fd;color:#1a5276}
.cb-caution{border-color:var(--accent);background:#fef9ec;color:#7d5a00}
.cb-warning{border-color:var(--red);background:#fdecea;color:#7b1a14}

/* Step */
.step-item{display:flex;gap:10px;padding:8px 0;
  border-bottom:1px solid #f0f3f8;font-size:13.5px;line-height:1.65;
  font-family:var(--mono);color:#2c3e50}
.step-item:last-child{border-bottom:none}
.step-num{color:var(--navy3);font-weight:700;flex-shrink:0;min-width:26px}

/* Image — inline in flow */
.inline-img-wrap{margin:16px 0;text-align:center}
.inline-img-wrap img{
  max-width:100%;border:1px solid var(--border);border-radius:8px;
  padding:6px;background:#fafafa;cursor:zoom-in;
  box-shadow:0 2px 10px rgba(15,42,74,.08);
  transition:box-shadow .15s}
.inline-img-wrap img:hover{box-shadow:0 4px 20px rgba(15,42,74,.18)}
.img-label{font-size:11px;color:var(--muted);font-family:var(--mono);
  margin-top:6px;display:flex;align-items:center;justify-content:center;gap:10px}
.img-dl{font-size:11px;color:var(--navy2);text-decoration:none;
  padding:2px 9px;border-radius:4px;border:1px solid var(--border)}
.img-dl:hover{background:#e8f0fe}

.tag-row{display:flex;flex-wrap:wrap;gap:6px;margin-top:6px}
.tag{display:inline-flex;align-items:center;gap:5px;padding:4px 12px;
  border-radius:20px;font-size:12px;font-family:var(--mono);cursor:default}
.tag-fig{background:#fef3e2;color:#7d5a00;border:1px solid #f5d8a0}
.tag-fig.fig-clickable:hover{background:#fde8b0;transform:translateY(-1px);box-shadow:0 2px 8px rgba(125,90,0,.2);transition:all .15s}
.tag-dc{background:#eafaf1;color:#1e7e34;border:1px solid #b7e8c5}

/* Lightbox */
.lightbox{display:none;position:fixed;inset:0;z-index:1000;
  background:rgba(10,20,40,.92);align-items:center;justify-content:center;padding:32px}
.lightbox.open{display:flex}
.lightbox img{max-width:90vw;max-height:90vh;border-radius:8px;
  box-shadow:0 20px 60px rgba(0,0,0,.5)}
.lb-close{position:absolute;top:20px;right:24px;color:white;font-size:32px;
  cursor:pointer;background:none;border:none;line-height:1;opacity:.8}
.lb-close:hover{opacity:1}

@media(max-width:600px){
  .srow{flex-wrap:wrap}
  .hstats{display:none}
  .content-flow{padding:20px 16px}
}
</style>
</head>
<body>

<header>
  <div class="hlogo">RR</div>
  <div>
    <div class="htitle">SAESL Subtask Lookup</div>
    <div class="hsub" id="mode-subtitle">Trent Engine Manual — Solumina Paste Tool</div>
  </div>
  <div class="hstats" id="hdr-stats"></div>
</header>

<div class="wrap">
  <div class="search-card">
    <span class="slabel">Enter Subtask ID</span>
    <div class="srow">
      <input type="text" id="sid-input"
             placeholder="e.g.  72-41-31-110-066-001"
             onkeydown="if(event.key==='Enter')fetchSubtask()">
      <button class="btn btn-primary" onclick="fetchSubtask()">🔍 Fetch</button>
      <button class="btn btn-success" id="paste-btn"
              onclick="pasteToTarget()" disabled>⬇ Paste</button>
      <button class="btn btn-ghost" onclick="clearResult()">Clear</button>
    </div>
    <div id="status" class="status"></div>
  </div>

  <div id="result-area">
    <div class="empty-state">
      <div class="empty-icon">📋</div>
      <p>Enter a subtask ID to load its content</p>
      <p style="font-size:13px;margin-top:8px;opacity:.7">
        e.g.&nbsp; 72-41-31-110-066-001 &nbsp;·&nbsp; press Enter or click Fetch
      </p>
    </div>
  </div>
</div>

<div class="lightbox" id="lightbox" onclick="closeLightbox()">
  <button class="lb-close">✕</button>
  <img id="lb-img" src="" alt="">
</div>

<script>
let currentData = null;
let outputMode  = 'solumina';  // updated from server on load

async function loadStats(){
  try{
    const d = await(await fetch('/api/stats')).json();
    document.getElementById('hdr-stats').innerHTML=
      pill(d.total_subtasks,'Subtasks')+pill(d.total_figures,'Figures');
  }catch(e){}
}
async function loadMode(){
  try{
    const d = await(await fetch('/api/mode')).json();
    outputMode = d.mode || 'solumina';
    const labels = {
      solumina: {sub:'Trent Engine Manual — Solumina Paste Tool', btn:'⬇ Paste to Solumina'},
      word:     {sub:'Trent Engine Manual — Open in WordPad',     btn:'📄 Open in WordPad'},
      notepad:  {sub:'Trent Engine Manual — Plain Text Copy',     btn:'📋 Copy Plain Text'},
    };
    const lbl = labels[outputMode] || labels.solumina;
    document.getElementById('mode-subtitle').textContent = lbl.sub;
    document.getElementById('paste-btn').textContent     = lbl.btn;
  }catch(e){}
}
function pill(v,l){return`<div class="hstat"><div class="val">${v.toLocaleString()}</div><div class="lbl">${l}</div></div>`}
loadStats();
loadMode();

async function fetchSubtask(){
  const id=document.getElementById('sid-input').value.trim();
  if(!id){showStatus('Please enter a subtask ID','error');return}
  showStatus('Loading…','info');
  document.getElementById('result-area').innerHTML=
    '<div class="loader"><div class="spinner"></div><span>Fetching subtask data…</span></div>';
  document.getElementById('paste-btn').disabled=true;
  try{
    const res=await fetch('/api/subtask/'+encodeURIComponent(id));
    const d=await res.json();
    if(!res.ok||d.error){
      showStatus(d.error||'Subtask not found','error');
      document.getElementById('result-area').innerHTML=
        '<div class="empty-state"><div class="empty-icon">❌</div>'
        +'<p>No subtask found for <strong>'+esc(id)+'</strong></p>'
        +'<p style="font-size:13px;margin-top:8px;opacity:.7">Ensure the PDF has been parsed.</p></div>';
      return;
    }
    currentData=d;
    renderResult(d);
    const imgs=(d.figures||[]).length;
    showStatus('✓ Loaded — '+d.steps_count+' steps · '+d.refs_count+' cross-refs'
      +(imgs?' · '+imgs+' image'+(imgs>1?'s':''):''),'success');
    document.getElementById('paste-btn').disabled=false;
  }catch(e){showStatus('Error: '+e.message,'error')}
}

function renderResult(d){
  const items=d.content_items||[];
  const imgs=(d.figures||[]).length;
  let stepNum=0;

  let html=`
    <div class="result-card">
      <div class="rh">
        <div>
          <div class="rh-id">${esc(d.subtask_id)}</div>
          ${d.title?'<div class="rh-title">'+esc(d.title)+'</div>':''}
        </div>
        <div class="rh-actions">
          <button class="btn btn-ghost"
            style="color:white;border-color:rgba(255,255,255,.35)"
            onclick="copyText()">📋 Copy</button>
        </div>
      </div>
      <div class="stat-strip">
        ${chip('📄','Task',d.task_id||'—')}
        ${chip('📅','Rev',d.revision_date||'—')}
        ${chip('🔢','Steps',d.steps_count)}
        ${chip('🖼','Images',imgs)}
      </div>
      <div class="content-flow">`;

  let inStepList=false;
  function closeStepList(){
    if(inStepList){html+='</div>';inStepList=false;}
  }

  const hasNotes   =(d.notes   ||[]).length>0;
  const hasCautions=(d.cautions||[]).length>0;
  const hasWarnings=(d.warnings||[]).length>0;
  let shownNoteHead=false, shownCautHead=false, shownWarnHead=false, shownStepHead=false;

  for(const ci of items){
    if(ci.type==='callout'){
      closeStepList();
      const low=(ci.text||'').toLowerCase();
      if(low.startsWith('caution')){
        if(!shownCautHead){html+='<div class="section-head">Cautions</div>';shownCautHead=true;}
        html+=`<div class="callout-box cb-caution">⚠️ ${esc(ci.text)}</div>`;
      } else if(low.startsWith('warning')||low.startsWith('notice')){
        if(!shownWarnHead){html+='<div class="section-head">Warnings</div>';shownWarnHead=true;}
        html+=`<div class="callout-box cb-warning">🚨 ${esc(ci.text)}</div>`;
      } else {
        if(!shownNoteHead){html+='<div class="section-head">Notes</div>';shownNoteHead=true;}
        html+=`<div class="callout-box cb-note">ℹ️ ${esc(ci.text)}</div>`;
      }

    } else if(ci.type==='image'){
      closeStepList();
      const src='/api/image?path='+encodeURIComponent(ci.path);
      const rawLabel=ci.label||'Image';
      const dispLabel=rawLabel.replace(/__p\d+__img\d+\.[^.]+$/, '').replace(/_/g,'-') || rawLabel;
      html+=`
        <div class="inline-img-wrap">
          <img src="${src}" alt="${esc(dispLabel)}" onclick="openLightbox('${src}')"
               onerror="this.closest('.inline-img-wrap').style.display='none'">
          <div class="img-label">
            <span>${esc(dispLabel)}</span>
            <a class="img-dl" href="${src}" download="${esc(rawLabel)}">↓ Save</a>
          </div>
        </div>`;

    } else if(ci.type==='step'||ci.type==='text'||ci.type==='row'){
      if(!shownStepHead&&(ci.type==='step')){
        closeStepList();
        html+='<div class="section-head">Procedure Steps</div>';
        shownStepHead=true;
      }
      if(!inStepList){
        html+='<div>';
        inStepList=true;
      }
      stepNum++;
      const clean=(ci.text||'').replace(/^\s*(?:[A-Z]\.|\(\d+\)|\([a-z]\)|\d{1,2}\.)\s+/,'');
      html+=`<div class="step-item">
        <span class="step-num">${String(stepNum).padStart(2,'0')}</span>
        <span>${esc(clean)}</span>
      </div>`;
    }
  }
  closeStepList();

  // ── Figure text refs (clickable, no cross-refs section) ───────────
  const figRefs=d.figure_refs||[];
  if(figRefs.length){
    html+='<div class="section-head" style="margin-top:20px">Figure References</div><div class="tag-row">';
    figRefs.forEach(f=>{html+=`<span class="tag tag-fig fig-clickable" onclick="loadFigure('${esc(f)}')" title="Click to view figure ${esc(f)}" style="cursor:pointer">📐 Fig ${esc(f)} <span style="font-size:10px;opacity:.7">🔍</span></span>`;});
    html+='</div>';
  }
  const dcs=d.data_cards||[];
  if(dcs.length){
    html+='<div class="section-head" style="margin-top:16px">Data Cards</div><div class="tag-row">';
    dcs.forEach(c=>{html+=`<span class="tag tag-dc">📎 ${esc(c)}</span>`;});
    html+='</div>';
  }

  html+='</div></div></div>';
  document.getElementById('result-area').innerHTML=html;
}

function chip(ico,lbl,val){
  return`<div class="stat-chip"><span>${ico}</span>${lbl}: <strong>${esc(String(val))}</strong></div>`;
}

async function loadFigure(figId){
  let modal=document.getElementById('fig-modal');
  if(!modal){
    modal=document.createElement('div');
    modal.id='fig-modal';
    modal.style.cssText='position:fixed;inset:0;background:rgba(0,0,0,.65);z-index:9999;display:flex;align-items:center;justify-content:center;padding:24px';
    modal.innerHTML=`
      <div id="fig-modal-box" style="background:#fff;border-radius:14px;max-width:860px;width:100%;max-height:90vh;overflow:auto;box-shadow:0 8px 48px rgba(0,0,0,.4)">
        <div style="display:flex;align-items:center;justify-content:space-between;padding:16px 24px;border-bottom:1px solid #dde3ef;position:sticky;top:0;background:#fff;z-index:1">
          <div>
            <div id="fig-modal-id" style="font-family:monospace;font-size:15px;font-weight:700;color:#0f2a4a"></div>
            <div id="fig-modal-caption" style="font-size:12px;color:#6b7a99;margin-top:2px"></div>
          </div>
          <button onclick="document.getElementById('fig-modal').remove()" style="border:none;background:#f0f3f8;border-radius:8px;padding:8px 14px;cursor:pointer;font-size:15px;color:#1c2535">✕ Close</button>
        </div>
        <div id="fig-modal-body" style="padding:24px;text-align:center">
          <div class="spinner" style="margin:40px auto"></div>
          <div style="color:#6b7a99;font-size:13px;margin-top:12px">Searching PDF for figure ${esc(figId)}…</div>
        </div>
      </div>`;
    modal.addEventListener('click',e=>{if(e.target===modal)modal.remove();});
    document.body.appendChild(modal);
  } else {
    document.getElementById('fig-modal-id').textContent='';
    document.getElementById('fig-modal-caption').textContent='';
    document.getElementById('fig-modal-body').innerHTML=`<div class="spinner" style="margin:40px auto"></div><div style="color:#6b7a99;font-size:13px;margin-top:12px">Searching PDF for figure ${esc(figId)}…</div>`;
    modal.style.display='flex';
  }
  document.getElementById('fig-modal-id').textContent='Fig '+figId;

  try{
    const resp=await fetch('/api/figure/'+encodeURIComponent(figId));
    if(!resp.ok){
      const err=await resp.json();
      document.getElementById('fig-modal-body').innerHTML=
        `<div style="padding:40px;color:#c0392b;font-size:14px">⚠️ ${esc(err.error||'Figure not found in index')}
        <br><br><span style="color:#6b7a99;font-size:12px">The figure image was not found in the indexed PDF. Make sure the PDF was parsed after this update.</span></div>`;
      return;
    }
    const fig=await resp.json();
    document.getElementById('fig-modal-caption').textContent=fig.caption||('Page '+(fig.page_num+1));
    document.getElementById('fig-modal-body').innerHTML=
      `<img src="${fig.image_url}" alt="Fig ${esc(fig.figure_id)}"
            style="max-width:100%;border:1px solid #dde3ef;border-radius:8px;box-shadow:0 2px 12px rgba(0,0,0,.1)"
            onerror="this.replaceWith(Object.assign(document.createElement('div'),{textContent:'⚠️ Image file missing on disk.',style:'color:#c0392b;padding:40px;font-size:14px'}))">
       <div style="font-size:11px;color:#6b7a99;margin-top:10px">Fig ${esc(fig.figure_id)} · page ${fig.page_num+1}</div>`;
  }catch(e){
    document.getElementById('fig-modal-body').innerHTML=`<div style="padding:40px;color:#c0392b">Network error: ${esc(String(e))}</div>`;
  }
}

async function copyText(){
  if(!currentData)return;
  try{
    await navigator.clipboard.writeText(currentData.formatted||'');
    showStatus('✓ Plain text copied to clipboard','success');
  }catch(e){showStatus('Copy failed: '+e.message,'error')}
}

async function pasteToTarget(){
  if(!currentData)return;
  if(outputMode==='notepad'){
    try{
      await navigator.clipboard.writeText(currentData.formatted||'');
      showStatus('✓ Plain text copied to clipboard — press Ctrl+V in Notepad','success');
    }catch(e){showStatus('Copy failed: '+e.message,'error')}
    return;
  }
  if(outputMode==='word'){
    showStatus('Building Word document with figures…','info');
    try{
      const res=await fetch('/api/paste_word',{
        method:'POST',
        headers:{'Content-Type':'application/json'},
        body:JSON.stringify({subtask_id:currentData.subtask_id})
      });
      const d=await res.json();
      if(d.success){
        showStatus('✓ Document created — WordPad is opening with all content and figures!','success');
      } else {
        showStatus('Failed: '+(d.error||'Unknown error'),'error');
      }
    }catch(e){showStatus('Error: '+e.message,'error')}
    return;
  }
  // solumina mode
  showStatus('Preparing rich HTML paste…','info');
  try{
    const res=await fetch('/api/paste',{
      method:'POST',
      headers:{'Content-Type':'application/json'},
      body:JSON.stringify({subtask_id:currentData.subtask_id})
    });
    const d=await res.json();
    if(d.success){showStatus('✓ Rich content (text + images) pasted into Solumina!','success');}
    else{showStatus('Paste failed: '+(d.error||'Unknown error'),'error');}
  }catch(e){showStatus('Error: '+e.message,'error')}
}

function clearResult(){
  document.getElementById('sid-input').value='';
  document.getElementById('result-area').innerHTML=
    '<div class="empty-state"><div class="empty-icon">📋</div>'
    +'<p>Enter a subtask ID to load its content</p></div>';
  document.getElementById('status').style.display='none';
  document.getElementById('paste-btn').disabled=true;
  currentData=null;
}

function showStatus(msg,type){
  const el=document.getElementById('status');
  el.textContent=msg; el.className='status '+type;
}

function esc(s){
  if(!s&&s!==0)return'';
  return String(s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;');
}

function openLightbox(src){
  document.getElementById('lb-img').src=src;
  document.getElementById('lightbox').classList.add('open');
}
function closeLightbox(){document.getElementById('lightbox').classList.remove('open');}
document.addEventListener('keydown',e=>{if(e.key==='Escape')closeLightbox();});
</script>
</body>
</html>
"""


# ─────────────────────────────────────────────
# Routes
# ─────────────────────────────────────────────

@app.route("/")
def index():
    return render_template_string(HTML)


@app.route("/api/subtask/<subtask_id>")
def api_subtask(subtask_id):
    db   = Database(DB_PATH)
    data = db.get_subtask_with_refs(subtask_id)
    db.close()

    if not data:
        return jsonify({"error": f"Subtask '{subtask_id}' not found"}), 404

    return jsonify({
        "subtask_id":      data["subtask_id"],
        "task_id":         data.get("task_id", ""),
        "title":           data.get("title", ""),
        "revision_date":   data.get("revision_date", ""),
        "formatted":       format_for_clipboard(data),
        "content_items":   data.get("content_items", []),
        "procedure_steps": data.get("procedure_steps", []),
        "notes":           data.get("notes", []),
        "cautions":        data.get("cautions", []),
        "warnings":        data.get("warnings", []),
        "steps_count":     len(data.get("procedure_steps", [])),
        "refs_count":      len(data.get("cross_refs", [])),
        "cross_refs":      data.get("cross_refs", []),
        "figure_refs":     data.get("figure_refs", []),
        "data_cards":      data.get("data_cards", []),
        "figures":         data.get("figures", []),
        "resolved_refs":   data.get("resolved_refs", []),
    })


@app.route("/api/image")
def api_image():
    path = request.args.get("path", "")
    if not path:
        abort(400)
    p = Path(path)
    if not p.exists() or not p.is_file():
        abort(404)
    try:
        p.resolve().relative_to(Path("output").resolve())
    except ValueError:
        abort(403)
    mime, _ = mimetypes.guess_type(str(p))
    return send_file(str(p), mimetype=mime or "image/png")


@app.route("/api/figure/<path:figure_id>")
def api_figure(figure_id):
    db  = Database(DB_PATH)
    fig = db.get_figure_image(figure_id)
    db.close()

    if not fig:
        return jsonify({"error": f"Figure '{figure_id}' not found in index"}), 404

    img_path = fig.get("image_path", "")
    if not img_path or not Path(img_path).exists():
        return jsonify({"error": "Figure image file not found on disk"}), 404

    return jsonify({
        "figure_id":  fig["figure_id"],
        "image_url":  f"/api/image?path={img_path}",
        "caption":    fig.get("caption", ""),
        "page_num":   fig.get("page_num", 0),
    })


@app.route("/api/paste_word", methods=["POST"])
def api_paste_word():
    """
    Builds a .docx with all subtask content + embedded figure images,
    writes it to a temp file, and opens it in WordPad automatically.
    No cross-references are included.
    """
    body       = request.get_json()
    subtask_id = body.get("subtask_id", "")

    db   = Database(DB_PATH)
    data = db.get_subtask_with_refs(subtask_id)
    db.close()

    if not data:
        return jsonify({"success": False, "error": "Subtask not found"})

    try:
        docx_bytes = build_docx(data)
    except Exception as e:
        return jsonify({"success": False, "error": f"Failed to build document: {e}"})

    # Write to a named temp file so WordPad can open it by path
    safe_id  = subtask_id.replace("/", "_").replace("\\", "_")
    tmp_path = Path(tempfile.gettempdir()) / f"saesl_{safe_id}.docx"
    tmp_path.write_bytes(docx_bytes)

    # Open in WordPad (Windows)
    try:
        wordpad = r"C:\Program Files\Windows NT\Accessories\wordpad.exe"
        if Path(wordpad).exists():
            subprocess.Popen([wordpad, str(tmp_path)])
        else:
            # Fallback: let Windows pick the default .docx handler
            os.startfile(str(tmp_path))
        return jsonify({"success": True})
    except Exception as e:
        return jsonify({"success": False, "error": f"Could not open WordPad: {e}"})


@app.route("/api/paste", methods=["POST"])
def api_paste():
    """
    Builds a rich HTML string with base64-embedded images and
    puts it on the clipboard via CF_HTML (for Solumina mode).
    """
    body       = request.get_json()
    subtask_id = body.get("subtask_id", "")

    db   = Database(DB_PATH)
    data = db.get_subtask_with_refs(subtask_id)
    db.close()

    if not data:
        return jsonify({"success": False, "error": "Subtask not found"})

    rich_html = build_rich_html(data, base_url="")

    full_html = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"></head>
<body style="font-family:'Segoe UI',Arial,sans-serif;color:#1c2535;padding:20px">
  <h2 style="font-family:Courier New,monospace;color:#1a3a6b">
    SUBTASK {data['subtask_id']}</h2>
  {f'<p style="font-style:italic;color:#6b7a99;margin:4px 0 16px">{_esc(data.get("title",""))}</p>' if data.get("title") else ""}
  {rich_html}
</body></html>"""

    try:
        result = _win32_html_clipboard(full_html)
        if result["success"]:
            if OUTPUT_MODE == "solumina":
                _send_paste_keystroke()
            return jsonify({"success": True})
    except Exception:
        pass

    try:
        pyperclip.copy(format_for_clipboard(data))
        return jsonify({
            "success": True,
            "warning": "Rich paste not available — plain text copied. Press Ctrl+V in Solumina."
        })
    except Exception as e:
        return jsonify({"success": False, "error": str(e)})


def _win32_html_clipboard(html: str) -> dict:
    try:
        import win32clipboard
        import win32con
    except ImportError:
        return {"success": False, "error": "win32clipboard not available"}

    header_template = (
        "Version:0.9\r\n"
        "StartHTML:{start_html:08d}\r\n"
        "EndHTML:{end_html:08d}\r\n"
        "StartFragment:{start_frag:08d}\r\n"
        "EndFragment:{end_frag:08d}\r\n"
    )
    FRAG_START = "<!--StartFragment-->"
    FRAG_END   = "<!--EndFragment-->"

    fragment = f"{FRAG_START}{html}{FRAG_END}"
    header_placeholder = header_template.format(
        start_html=0, end_html=0, start_frag=0, end_frag=0)
    offset = len(header_placeholder.encode("utf-8"))
    start_html = offset
    start_frag = offset + len(FRAG_START.encode("utf-8"))
    end_frag   = start_frag + len(html.encode("utf-8"))
    end_html   = end_frag + len(FRAG_END.encode("utf-8"))

    header = header_template.format(
        start_html=start_html, end_html=end_html,
        start_frag=start_frag, end_frag=end_frag,
    )
    cf_html = (header + fragment).encode("utf-8")

    try:
        win32clipboard.OpenClipboard()
        win32clipboard.EmptyClipboard()
        cf_html_fmt = win32clipboard.RegisterClipboardFormat("HTML Format")
        win32clipboard.SetClipboardData(cf_html_fmt, cf_html)
        win32clipboard.SetClipboardData(win32con.CF_UNICODETEXT, html)
        win32clipboard.CloseClipboard()
        return {"success": True}
    except Exception as e:
        try:
            win32clipboard.CloseClipboard()
        except Exception:
            pass
        return {"success": False, "error": str(e)}


def _send_paste_keystroke():
    try:
        import time
        import win32api, win32con, win32gui
        hwnd = win32gui.GetForegroundWindow()
        if hwnd:
            win32api.keybd_event(win32con.VK_CONTROL, 0, 0, 0)
            win32api.keybd_event(ord('V'), 0, 0, 0)
            win32api.keybd_event(ord('V'), 0, win32con.KEYEVENTF_KEYUP, 0)
            win32api.keybd_event(win32con.VK_CONTROL, 0, win32con.KEYEVENTF_KEYUP, 0)
            time.sleep(0.1)
    except Exception:
        pass


@app.route("/api/stats")
def api_stats():
    db    = Database(DB_PATH)
    stats = db.stats()
    db.close()
    return jsonify(stats)


@app.route("/api/mode")
def api_mode():
    return jsonify({"mode": OUTPUT_MODE})


@app.route("/api/search")
def api_search():
    db = Database(DB_PATH)
    results = db.search_subtasks(request.args.get("q", ""))
    db.close()
    return jsonify(results)


# ─────────────────────────────────────────────
# Entry point
# ─────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--db",   default="output/saesl.db")
    parser.add_argument("--port", default=5000, type=int)
    parser.add_argument("--host", default="127.0.0.1")
    args = parser.parse_args()
    DB_PATH = args.db
    print(f"\n{'='*50}\n  SAESL Subtask Lookup\n  DB: {DB_PATH}\n  URL: http://{args.host}:{args.port}\n{'='*50}\n")
    app.run(host=args.host, port=args.port, debug=False)
